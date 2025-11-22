import pandas as pd
from docx import Document
from tkinter import Tk, Label, Entry, Button, filedialog, messagebox

def create_test_report(requirement_item, formal_content, test_case_path, report_file):
    # 创建 Word 文档
    doc = Document()
    doc.add_heading('测试报告', level=1)

    # 添加试验需求
    doc.add_heading('1.1.1 试验需求', level=2)
    doc.add_paragraph(requirement_item)

    # 添加试验目的
    doc.add_heading('1.1.2 试验目的', level=2)
    doc.add_paragraph(f"验证需求：[{requirement_item}]")

    # 添加试验步骤
    doc.add_heading('1.1.3 试验步骤', level=2)
    doc.add_paragraph(f"1 将需求条目形式化如下：【{formal_content}】")
    doc.add_paragraph("2 生成测试用例表并方针执行，结果如下：")

    # 读取测试用例文件
    try:
        test_cases = pd.read_excel(test_case_path)
    except Exception as e:
        messagebox.showerror("错误", f"加载测试用例表格失败：{e}")
        return

    # 创建表格并添加数据
    table = doc.add_table(rows=1, cols=len(test_cases.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(test_cases.columns):
        hdr_cells[i].text = column_name

    for idx, row in test_cases.iterrows():
        row_cells = table.add_row().cells
        for i, column_name in enumerate(test_cases.columns):
            row_cells[i].text = str(row[column_name])

    # 添加实验结果
    doc.add_heading('1.1.4 实验结果', level=2)
    test_result = check_test_case_results(test_cases)
    doc.add_paragraph(test_result)

    try:
        doc.save(report_file)
        messagebox.showinfo("成功", f"测试报告已生成：{report_file}")
    except Exception as e:
        messagebox.showerror("错误", f"保存测试报告失败：{e}")

def check_test_case_results(test_cases):
    last_two_columns = test_cases.iloc[:, -2:]
    results = []

    for idx in range(len(last_two_columns)):
        expected_output = str(last_two_columns.iloc[idx, 0])  # 转换为字符串
        model_output = str(last_two_columns.iloc[idx, 1])  # 转换为字符串

        # 检查是否包含字符串“不小于”或“小于”
        if "不小于" in expected_output:
            # 提取期望值并进行数值比较
            threshold = float(expected_output.replace("不小于", "").strip())
            try:
                model_output_value = float(model_output)
                results.append(model_output_value >= threshold)
            except ValueError:
                results.append(False)
        elif "小于" in expected_output:
            # 提取期望值并进行数值比较
            threshold = float(expected_output.replace("小于", "").strip())
            try:
                model_output_value = float(model_output)
                results.append(model_output_value < threshold)
            except ValueError:
                results.append(False)
        else:
            # 直接进行字符串比较
            results.append(expected_output == model_output)

    if all(results):
        return "测试通过：所有模型输出符合期望输出内容。"
    else:
        return "测试不通过：存在不一致的测试用例。"

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    test_case_path_entry.delete(0, 'end')
    test_case_path_entry.insert(0, file_path)

def generate_report():
    requirement_item = requirement_item_entry.get()
    formal_content = formal_content_entry.get()
    test_case_path = test_case_path_entry.get()
    report_file = report_file_entry.get()

    if not requirement_item or not formal_content or not test_case_path or not report_file:
        messagebox.showwarning("警告", "请填写所有字段！")
    else:
        create_test_report(requirement_item, formal_content, test_case_path, report_file)

# 设置 Tkinter 窗口
root = Tk()
root.title("测试报告生成器")
root.geometry("400x300")

Label(root, text="需求条目：").pack(pady=5)
requirement_item_entry = Entry(root, width=50)
requirement_item_entry.pack()

Label(root, text="形式化内容：").pack(pady=5)
formal_content_entry = Entry(root, width=50)
formal_content_entry.pack()

Label(root, text="测试用例文件路径：").pack(pady=5)
test_case_path_entry = Entry(root, width=50)
test_case_path_entry.pack()
Button(root, text="选择文件", command=select_file).pack()

Label(root, text="报告文件名：").pack(pady=5)
report_file_entry = Entry(root, width=50)
report_file_entry.insert(0, "测试报告.docx")  # 设置默认文件名
report_file_entry.pack()

Button(root, text="生成报告", command=generate_report).pack(pady=20)

root.mainloop()
